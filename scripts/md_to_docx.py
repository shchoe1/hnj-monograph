#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
md_to_docx.py · 챕터 Markdown -> Word(.docx) 변환 (검토용)
힘뇌장(HIMNOEJANG) 모노그래프 · 산출 도구

이 저장소 챕터의 마크다운 관례를 처리한다:
  # / ## / ### 제목, **굵게**, `코드`, > 인용블록, --- 수평선,
  | 표 |, - 불릿, 1. 번호목록, [text](url)/<url> 링크(텍스트로 평문화).

한글이 깨지지 않도록 기본 글꼴을 맑은 고딕(Malgun Gothic)으로 지정한다.
의존: python-docx (pip install python-docx)

사용법:
  python scripts/md_to_docx.py chapters/ch12-glutamine.md            # -> output/ch12-glutamine.docx
  python scripts/md_to_docx.py chapters/ch12-glutamine.md out.docx   # 명시 경로
  python scripts/md_to_docx.py --all                                 # chapters/ch*.md 전부 -> output/
  python scripts/md_to_docx.py --merge output/HIMNOEJANG_Ch1-12.docx # chapters/ch*.md 전부 -> 통합본 1개
"""
import sys, os, re, glob
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KFONT = "Malgun Gothic"  # 한글 렌더링


def set_korean_font(doc):
    style = doc.styles["Normal"]
    style.font.name = KFONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(a), KFONT)


def _apply_font(run):
    run.font.name = KFONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KFONT)


# 인라인: **굵게**, `코드`, [txt](url), <url>
LINK_MD = re.compile(r'\[([^\]]+)\]\((?:[^)]+)\)')
AUTOLINK = re.compile(r'<((?:https?://|mailto:)[^>]+)>')
TOKEN = re.compile(r'(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+`)')


def add_runs(paragraph, text):
    text = LINK_MD.sub(r'\1', text)        # [txt](url) -> txt
    text = AUTOLINK.sub(r'\1', text)       # <url> -> url
    text = text.replace("\\&", "&").replace("\\%", "%").replace("\\_", "_")
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"
        else:
            r = paragraph.add_run(part)
        _apply_font(r)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom); pPr.append(pbdr)


def add_table(doc, rows):
    # rows: 리스트[리스트[셀문자열]] (구분선 행 제외)
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        for j in range(ncol):
            txt = r[j] if j < len(r) else ""
            para = cells[j].paragraphs[0]
            add_runs(para, txt)
            if i == 0:
                for run in para.runs:
                    run.bold = True


def is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def is_sep_row(line):
    return bool(re.match(r'^\s*\|[\s:|-]+\|\s*$', line))


def split_cells(line):
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def render(doc, lines):
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1; continue

        # 표 블록
        if is_table_row(line):
            block = []
            while i < n and is_table_row(lines[i]):
                if not is_sep_row(lines[i]):
                    block.append(split_cells(lines[i]))
                i += 1
            if block:
                add_table(doc, block)
            continue

        # 수평선
        if re.match(r'^\s*---+\s*$', line) or re.match(r'^\s*\*\*\*+\s*$', line):
            add_hr(doc); i += 1; continue

        # 제목
        m = re.match(r'^(#{1,6})\s+(.*)$', s)
        if m:
            level = len(m.group(1))
            p = doc.add_heading(level=min(level, 4))
            p.text = ""
            add_runs(p, m.group(2))
            i += 1; continue

        # 인용블록 (연속 > 라인)
        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(re.sub(r'^\s*>\s?', '', lines[i]))
                i += 1
            p = doc.add_paragraph(style="Intense Quote")
            p.text = ""
            add_runs(p, " ".join(x.strip() for x in buf if x.strip()))
            continue

        # 번호 목록
        m = re.match(r'^(\d+)\.\s+(.*)$', s)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.text = ""
            add_runs(p, m.group(2))
            i += 1; continue

        # 불릿 목록
        m = re.match(r'^[-*]\s+(.*)$', s)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.text = ""
            add_runs(p, m.group(1))
            i += 1; continue

        # 일반 문단
        p = doc.add_paragraph()
        add_runs(p, s)
        i += 1


def convert(md_path, out_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")
    doc = Document()
    set_korean_font(doc)
    render(doc, lines)
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


def merge(md_files, out_path):
    from docx.enum.text import WD_BREAK
    doc = Document()
    set_korean_font(doc)
    for idx, md in enumerate(md_files):
        if idx > 0:
            doc.add_page_break()
        with open(md, encoding="utf-8") as f:
            render(doc, f.read().split("\n"))
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    args = sys.argv[1:]
    if not args:
        print(__doc__); sys.exit(1)
    if args[0] == "--all":
        files = sorted(glob.glob("chapters/ch*.md"))
        files = [f for f in files if os.path.basename(f) != "test.md"]
        for f in files:
            out = os.path.join("output", os.path.splitext(os.path.basename(f))[0] + ".docx")
            convert(f, out)
            print(f"  {f}  ->  {out}")
        print(f"\n변환 완료: {len(files)}개")
    elif args[0] == "--merge":
        out = args[1] if len(args) > 1 else os.path.join("output", "HIMNOEJANG_Ch1-12.docx")
        files = sorted(glob.glob("chapters/ch*.md"))
        files = [f for f in files if os.path.basename(f) != "test.md"]
        merge(files, out)
        print(f"통합본 {len(files)}개 챕터  ->  {out}")
    else:
        md = args[0]
        out = args[1] if len(args) > 1 else os.path.join(
            "output", os.path.splitext(os.path.basename(md))[0] + ".docx")
        convert(md, out)
        print(f"{md}  ->  {out}")


if __name__ == "__main__":
    main()
